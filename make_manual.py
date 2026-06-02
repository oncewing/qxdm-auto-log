"""
QXDM AUTO LOG 사용자 매뉴얼 생성기
실행  : python make_manual.py
결과  : QXDM_AUTO_LOG_Manual_v{version}.docx
의존성: pip install python-docx matplotlib
"""

import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    from matplotlib.patches import FancyBboxPatch
    HAS_MPL = True
except ImportError:
    HAS_MPL = False
    print("[경고] matplotlib 없음 → pip install matplotlib  (흐름도 이미지 생략)")

__version__ = "1.0.0"
APP_NAME    = "QXDM AUTO LOG"
OUT_FILE    = f"QXDM_AUTO_LOG_Manual_v{__version__}.docx"

COLOR_TITLE  = RGBColor(0x1F, 0x49, 0x7D)
COLOR_HEAD   = RGBColor(0x2E, 0x74, 0xB5)
COLOR_GRAY   = RGBColor(0x40, 0x40, 0x40)
COLOR_CODE_BG = "D6DCE4"   # 코드 블록 배경 (16진)
COLOR_TH_BG   = "2E74B5"   # 표 헤더 배경


# ─────────────────────────────────────────────
# XML 헬퍼
# ─────────────────────────────────────────────
def _set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def _set_cell_border(cell, **edges):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, val in edges.items():
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"),   val.get("val",   "single"))
        e.set(qn("w:sz"),    val.get("sz",    "4"))
        e.set(qn("w:space"), val.get("space", "0"))
        e.set(qn("w:color"), val.get("color", "auto"))
        tcBorders.append(e)
    tcPr.append(tcBorders)

def _para_space(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:before"), str(before))
    spc.set(qn("w:after"),  str(after))
    pPr.append(spc)


# ─────────────────────────────────────────────
# 공통 스타일 적용
# ─────────────────────────────────────────────
def _apply_default_styles(doc):
    style = doc.styles["Normal"]
    font  = style.font
    font.name = "맑은 고딕"
    font.size = Pt(10)

    for hname, sz, bold in [
        ("Heading 1", 14, True),
        ("Heading 2", 12, True),
        ("Heading 3", 11, True),
    ]:
        h = doc.styles[hname]
        h.font.name  = "맑은 고딕"
        h.font.size  = Pt(sz)
        h.font.bold  = bold
        h.font.color.rgb = COLOR_HEAD


# ─────────────────────────────────────────────
# 제목 페이지
# ─────────────────────────────────────────────
def _title_page(doc):
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(APP_NAME)
    r.font.size = Pt(32)
    r.font.bold = True
    r.font.color.rgb = COLOR_TITLE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"사용자 매뉴얼  v{__version__}")
    r2.font.size = Pt(16)
    r2.font.color.rgb = COLOR_HEAD

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(
        "QXDM을 자동으로 제어하여 모뎀 상태(LPM ↔ ONLINE) 사이클을 반복하고,\n"
        "사용자 정의 조건이 만족되지 않을 때 QXDM 로그를 자동 저장하는 GUI 도구"
    )
    r3.font.size = Pt(11)
    r3.font.color.rgb = COLOR_GRAY

    for _ in range(6):
        doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("내부 사용 도구  |  2026")
    r4.font.size = Pt(10)
    r4.font.color.rgb = COLOR_GRAY

    doc.add_page_break()


# ─────────────────────────────────────────────
# 표 헬퍼
# ─────────────────────────────────────────────
def _add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 헤더 행
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        _set_cell_bg(hdr_cells[i], COLOR_TH_BG)
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name  = "맑은 고딕"
        run.font.size  = Pt(10)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 데이터 행
    for ri, row in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    run.font.name = "맑은 고딕"
                    run.font.size = Pt(10)
            if ri % 2 == 1:
                _set_cell_bg(cells[ci], "EBF3FB")

    # 열 너비
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return tbl


# ─────────────────────────────────────────────
# 코드 블록
# ─────────────────────────────────────────────
def _add_code(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, COLOR_CODE_BG)
    cell.text = ""
    for line in text.splitlines():
        p = cell.add_paragraph(line)
        _para_space(p, before=0, after=0)
        for run in p.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(9)
    doc.add_paragraph()


# ─────────────────────────────────────────────
# 참고 박스 (배경색 단락)
# ─────────────────────────────────────────────
def _add_note(doc, text, bg="FFF2CC"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, bg)
    cell.text = ""
    p = cell.add_paragraph(f"  ⚠  {text}")
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt(10)
    doc.add_paragraph()


# ─────────────────────────────────────────────
# GUI 화면 placeholder
# ─────────────────────────────────────────────
def _add_screenshot_placeholder(doc, caption="GUI 실행 화면"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, "D9D9D9")
    cell.height = Cm(8)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[ {caption} ]\n실제 실행 화면 캡처 이미지를 삽입하세요.")
    r.font.name  = "맑은 고딕"
    r.font.size  = Pt(11)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    p2 = doc.add_paragraph(caption)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p2.runs:
        run.font.name  = "맑은 고딕"
        run.font.size  = Pt(9)
        run.font.italic = True
        run.font.color.rgb = COLOR_GRAY
    doc.add_paragraph()


# ─────────────────────────────────────────────
# matplotlib 흐름도 이미지
# ─────────────────────────────────────────────
def _flow_image_test() -> io.BytesIO | None:
    if not HAS_MPL:
        return None
    fig, ax = plt.subplots(figsize=(5, 8))
    ax.set_xlim(0, 10); ax.set_ylim(0, 20); ax.axis("off")

    def box(x, y, w, h, text, fc="#2E74B5", tc="white", fs=9, style="round,pad=0.1"):
        ax.add_patch(FancyBboxPatch((x, y), w, h,
                     boxstyle=style, fc=fc, ec="#1F497D", lw=1.2))
        ax.text(x + w/2, y + h/2, text, ha="center", va="center",
                fontsize=fs, color=tc, fontfamily="sans-serif",
                fontweight="bold", wrap=True,
                multialignment="center")

    def arr(x, y1, y2):
        ax.annotate("", xy=(x, y2), xytext=(x, y1),
                    arrowprops=dict(arrowstyle="->", color="#333", lw=1.5))

    def diamond(x, y, w, h, text, fc="#ED7D31"):
        xs = [x+w/2, x+w, x+w/2, x, x+w/2]
        ys = [y+h,   y+h/2, y,   y+h/2, y+h]
        ax.fill(xs, ys, fc=fc, ec="#C55A11", lw=1.2)
        ax.text(x+w/2, y+h/2, text, ha="center", va="center",
                fontsize=8, color="white", fontfamily="sans-serif",
                fontweight="bold", multialignment="center")

    # 박스들
    box(3, 18.2, 4, 1,   "▶ START",             fc="#1F497D")
    arr(5, 18.2, 17.5)
    box(2.5, 16.5, 5, 1, "adb 기기 확인")
    arr(5, 16.5, 15.8)
    diamond(3, 14.5, 4, 1.2, "강제 Crash\n옵션 ON?", fc="#ED7D31")
    # Yes 화살표
    ax.annotate("", xy=(5, 13.5), xytext=(5, 14.5),
                arrowprops=dict(arrowstyle="->", color="#333", lw=1.5))
    ax.text(5.2, 14.0, "Yes", fontsize=8, color="#333")
    box(2.5, 12.3, 5, 1.1, "AT!ERROROPT? 확인\n→ USB Dump 없으면 재부팅")
    arr(5, 12.3, 11.5)
    box(2.5, 10.3, 5, 1.1, "단말 복귀 대기\n(최대 120초)")
    # No 화살표 (우측 우회)
    ax.annotate("", xy=(8, 10.8), xytext=(7, 15.1),
                arrowprops=dict(arrowstyle="->", color="#888", lw=1.2,
                                connectionstyle="arc3,rad=-0.4"))
    ax.text(8.2, 13.0, "No", fontsize=8, color="#888")
    ax.annotate("", xy=(5, 9.5), xytext=(8, 10.8),
                arrowprops=dict(arrowstyle="->", color="#888", lw=1.2))

    arr(5, 10.3, 9.3)
    box(2.5, 8.1, 5, 1.1, "QXDM 실행")
    arr(5, 8.1, 7.2)
    box(2.5, 6.0, 5, 1.1, "Phone 연결 (Ctrl+O)\n[✔ 연결 완료] 클릭")
    arr(5, 6.0, 5.1)
    box(2.5, 3.8, 5, 1.1, "테스트 사이클 반복\n(LPM ↔ ONLINE)")
    arr(5, 3.8, 2.9)
    box(2.5, 1.6, 5, 1.1, "정상 조건 체크\n비정상 시 ISF 저장 후 종료", fc="#C00000")

    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _flow_image_crash() -> io.BytesIO | None:
    if not HAS_MPL:
        return None
    fig, ax = plt.subplots(figsize=(6, 7))
    ax.set_xlim(0, 12); ax.set_ylim(0, 14); ax.axis("off")

    def box(x, y, w, h, text, fc="#2E74B5", tc="white", fs=9):
        ax.add_patch(FancyBboxPatch((x, y), w, h,
                     boxstyle="round,pad=0.1", fc=fc, ec="#1F497D", lw=1.2))
        ax.text(x+w/2, y+h/2, text, ha="center", va="center",
                fontsize=fs, color=tc, fontweight="bold",
                multialignment="center")

    def diamond(x, y, w, h, text, fc="#ED7D31"):
        xs = [x+w/2, x+w, x+w/2, x, x+w/2]
        ys = [y+h,   y+h/2, y,   y+h/2, y+h]
        ax.fill(xs, ys, fc=fc, ec="#C55A11", lw=1.2)
        ax.text(x+w/2, y+h/2, text, ha="center", va="center",
                fontsize=8, color="white", fontweight="bold",
                multialignment="center")

    def arr(x1, y1, x2, y2, label="", lx=None, ly=None):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="->", color="#333", lw=1.5))
        if label:
            ax.text(lx or (x1+x2)/2+0.15, ly or (y1+y2)/2,
                    label, fontsize=8, color="#333")

    box(3.5, 12.5, 5, 1.0, "테스트 시작 (강제 Crash ON)", fc="#1F497D")
    arr(6, 12.5, 6, 12.0)
    box(3.5, 10.8, 5, 1.1, "AT!ERROROPT? 확인")
    arr(6, 10.8, 6, 10.2)
    diamond(3.5, 8.8, 5, 1.3, "응답에\nUSB Dump 포함?", fc="#ED7D31")

    # Yes → right
    ax.annotate("", xy=(9.5, 9.45), xytext=(8.5, 9.45),
                arrowprops=dict(arrowstyle="->", color="#333", lw=1.5))
    ax.text(8.55, 9.6, "포함", fontsize=8, color="#333")
    box(9.5, 8.95, 2.3, 1.0, "재부팅 없이\n진행", fc="#70AD47", fs=8)

    # No → down
    arr(6, 8.8, 6, 8.0, label="미포함", lx=6.1, ly=8.4)
    box(3.5, 6.8, 5, 1.1, "AT!ERROROPT=0\nAT+CFUN=1,1 (재부팅)")
    arr(6, 6.8, 6, 6.1)
    box(3.5, 4.9, 5, 1.1, "adb devices 폴링\n(최대 120초)")
    arr(6, 4.9, 6, 4.2)
    diamond(3.5, 3.0, 5, 1.1, "단말\n복귀?", fc="#ED7D31")

    # 복귀 → down
    arr(6, 3.0, 6, 2.2, label="복귀", lx=6.1, ly=2.6)
    box(3.5, 1.0, 5, 1.1, "QXDM 실행 → 테스트 시작", fc="#70AD47")

    # 미복귀 → right
    ax.annotate("", xy=(9.5, 3.55), xytext=(8.5, 3.55),
                arrowprops=dict(arrowstyle="->", color="#C00000", lw=1.5))
    ax.text(8.55, 3.7, "미복귀", fontsize=8, color="#C00000")
    box(9.5, 3.05, 2.3, 1.0, "테스트\n중단", fc="#C00000", fs=8)

    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _add_flow_image(doc, buf, caption, width=Inches(3.5)):
    if buf is None:
        _add_code(doc, f"[흐름도: {caption}]\nmatplotlib 설치 후 재실행하면 이미지로 표시됩니다.")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=width)
    p2 = doc.add_paragraph(caption)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p2.runs:
        r.font.name   = "맑은 고딕"
        r.font.size   = Pt(9)
        r.font.italic = True
        r.font.color.rgb = COLOR_GRAY
    doc.add_paragraph()


# ─────────────────────────────────────────────
# 문서 섹션 작성
# ─────────────────────────────────────────────
def _section_overview(doc):
    doc.add_heading("1. 프로그램 개요", level=1)
    p = doc.add_paragraph(
        f"{APP_NAME}는 QXDM을 자동으로 제어하여 모뎀 상태(LPM ↔ ONLINE) 사이클을 반복하고, "
        "사용자가 정의한 정상 조건이 만족되지 않을 때 QXDM 로그를 자동 저장하는 GUI 도구입니다."
    )
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt(10)

    doc.add_heading("주요 기능", level=2)
    features = [
        "ADB를 통한 모뎀 AT 명령 / Shell 명령 자동 실행",
        "사용자 정의 테스트 단계 구성 (AT / SHELL)",
        "사용자 정의 정상 조건 감시 (SHELL / AT / Python 함수)",
        "QXDM ISF 로그 자동 저장 및 롤링 관리",
        "조건 비정상 감지 시 QXDM 자동 종료 및 로그 보존",
        "트리거 발생 시 강제 Crash (sysrq-trigger) 지원",
        "테스트 시작 전 AT!ERROROPT 자동 확인 및 설정",
    ]
    for feat in features:
        p = doc.add_paragraph(feat, style="List Bullet")
        for run in p.runs:
            run.font.name = "맑은 고딕"
            run.font.size = Pt(10)
    doc.add_paragraph()


def _section_requirements(doc):
    doc.add_heading("2. 사전 요구 사항", level=1)
    _add_table(doc,
        headers=["항목", "내용"],
        rows=[
            ["Windows OS",    "Windows 10 이상 권장"],
            ["QXDM 설치",     "COM 자동화 인터페이스 필요 (QXDM5AutoApplication)"],
            ["ADB 설치",      "PATH에 adb.exe 등록 필요"],
            ["단말 USB 연결", "ADB 디버깅 활성화 필요"],
            ["Python",        "별도 설치 불필요 (EXE 내 포함)"],
        ],
        col_widths=[4, 11],
    )


def _section_screen(doc):
    doc.add_heading("3. 화면 구성", level=1)
    _add_screenshot_placeholder(doc, "QXDM AUTO LOG 메인 화면")

    doc.add_heading("영역 구분", level=2)
    _add_table(doc,
        headers=["영역", "설명"],
        rows=[
            ["테스트 파라미터", "Device, 저장 폴더, 사이클 수, 롤링/Crash 옵션 설정"],
            ["테스트 단계",     "LPM→ONLINE 순서로 실행될 AT/Shell 명령 목록"],
            ["정상 조건",       "마지막 단계 후 체크할 조건 목록 (AND/OR 로직)"],
            ["컨트롤",          "START / STOP / Phone 연결 완료 버튼, 상태 표시"],
            ["로그 출력",       "실시간 실행 로그 (색상 구분, 파일 동시 저장)"],
        ],
        col_widths=[3.5, 11.5],
    )


def _section_params(doc):
    doc.add_heading("4. 테스트 파라미터", level=1)
    _add_table(doc,
        headers=["항목", "설명", "기본값"],
        rows=[
            ["Device",              "ADB 연결 단말 선택. ↺ 새로고침으로 목록 갱신",                  "자동 감지"],
            ["저장 폴더",           "ISF 로그 및 텍스트 로그가 저장될 경로",                          r"D:\QXDM_LOGS"],
            ["사이클 수",           "LPM → ONLINE 반복 최대 횟수",                                    "100"],
            ["트리거 후 수집(초)",  "비정상 감지 후 추가 로그 수집 시간",                             "5"],
            ["롤링 저장 주기",      "N 사이클마다 ISF 저장 후 클리어. 0 = 비활성",                   "0"],
            ["롤링 보관 개수",      "롤링 저장 시 유지할 최대 파일 수 (초과 시 오래된 것 삭제)",      "5"],
            ["트리거 발생 시\n강제 Crash",
             "시작 전 AT!ERROROPT 확인 → USB Dump 없으면 AT!ERROROPT=0 + AT+CFUN=1,1 재부팅.\n"
             "트리거 발생 시 sysrq-trigger 실행", "OFF"],
            ["테스트 후\nWindows 종료", "테스트 완료 후 PC 자동 종료",                               "OFF"],
            ["종료 대기(초)",       "Windows 종료까지 대기 시간. shutdown /a 로 취소 가능",           "60"],
        ],
        col_widths=[3.5, 8.5, 3],
    )


def _section_steps(doc):
    doc.add_heading("5. 테스트 단계", level=1)
    p = doc.add_paragraph(
        "ONLINE 전환 전 순서대로 실행되는 명령 목록입니다. "
        "기본 구성은 LPM → ONLINE 두 단계이며, 자유롭게 추가/편집/삭제할 수 있습니다."
    )
    for run in p.runs:
        run.font.name = "맑은 고딕"; run.font.size = Pt(10)

    doc.add_heading("단계 편집 항목", level=2)
    _add_table(doc,
        headers=["항목", "설명"],
        rows=[
            ["라벨",             "QXDM 어노테이션에 기록될 이름 (예: LPM, ONLINE)"],
            ["타입",             "AT 또는 SHELL"],
            ["명령",             "AT 타입: AT+CFUN=4 형식 / SHELL 타입: adb shell 명령"],
            ["대기 기본(초)",    "명령 실행 후 고정 대기 시간"],
            ["랜덤 추가 최소/최대(초)", "기본 + random(최소~최대) 초 대기"],
        ],
        col_widths=[4, 11],
    )

    doc.add_heading("기본 단계 구성", level=2)
    _add_code(doc,
        "#1  [AT]  AT+CFUN=4   →  LPM 전환 후 7 + rand(1~15)초 대기\n"
        "#2  [AT]  AT+CFUN=1   →  ONLINE 전환 후 10 + rand(5~20)초 대기"
    )
    _add_note(doc, "AT 명령은 내부적으로 /data/shellat7 'AT+...' 로 실행됩니다.")


def _section_conditions(doc):
    doc.add_heading("6. 정상 조건", level=1)
    p = doc.add_paragraph(
        "마지막 테스트 단계 완료 후 조건을 체크합니다. "
        "조건이 비정상(False)이면 QXDM 로그를 저장하고 종료합니다."
    )
    for run in p.runs:
        run.font.name = "맑은 고딕"; run.font.size = Pt(10)

    doc.add_heading("조건 로직", level=2)
    _add_table(doc,
        headers=["로직", "동작"],
        rows=[
            ["AND", "모든 조건이 정상이어야 테스트 계속. 하나라도 비정상 → QXDM 종료"],
            ["OR",  "조건 중 하나만 정상이어도 테스트 계속. 전부 비정상 → QXDM 종료"],
        ],
        col_widths=[2.5, 12.5],
    )

    doc.add_heading("SHELL 타입", level=2)
    _add_table(doc,
        headers=["항목", "설명"],
        rows=[
            ["Shell 명령", "실행할 shell 명령 (예: ifconfig | grep 192.0.0.2)"],
            ["기대값",     "출력에서 찾을 문자열 (예: 192.0.0.2)"],
            ["정상 조건",  "출현 = 정상 : 기대값이 출현하면 정상\n미출현 = 정상 : 기대값이 없으면 정상"],
            ["타임아웃",   "조건 만족 대기 최대 시간 (초)"],
        ],
        col_widths=[3.5, 11.5],
    )

    doc.add_heading("AT 타입", level=2)
    _add_table(doc,
        headers=["항목", "설명"],
        rows=[
            ["AT 명령", "전송할 AT 명령 (예: AT+CFUN?)"],
            ["기대값",  "응답에서 찾을 문자열 (예: +CFUN: 1)"],
            ["정상 조건", "SHELL 타입과 동일"],
            ["타임아웃",  "응답 대기 최대 시간 (초)"],
        ],
        col_widths=[3.5, 11.5],
    )

    doc.add_heading("FUNCTION 타입", level=2)
    p = doc.add_paragraph("Python 코드를 직접 작성하여 복잡한 조건을 구현합니다.")
    for run in p.runs:
        run.font.name = "맑은 고딕"; run.font.size = Pt(10)
    _add_code(doc,
        "# 반환: (triggered: bool, reason: str)\n"
        "#   True  = 비정상 감지 → QXDM 종료\n"
        "#   False = 정상         → 테스트 계속\n\n"
        "# 사용 가능한 함수\n"
        "adb_shell('shell 명령')  # ADB shell 실행, stdout 반환\n"
        "adb_at('AT 명령')        # AT 명령 전송, 응답 반환\n\n"
        "# 작성 예시\n"
        "out = adb_shell('cat /var/tmp/test_func 2>/dev/null | grep ok')\n"
        "if out:\n"
        "    return False, 'ok_found'    # ok 확인됨 = 정상\n"
        "return True, 'ok_not_found'     # ok 없음 = 비정상"
    )
    _add_note(doc, "파일/명령 오류 메시지가 결과에 섞이지 않도록 2>/dev/null 사용을 권장합니다.")

    doc.add_heading("기본 조건 구성", level=2)
    _add_code(doc,
        "조건 1  [SHELL]  ifconfig | grep rmnet_data  →  'rmnet_data'  [출현=정상]  (15s)\n"
        "조건 2  [SHELL]  ifconfig | grep 192.0.0.2   →  '192.0.0.2'   [출현=정상]  (10s)"
    )


def _section_procedure(doc):
    doc.add_heading("7. 테스트 실행 절차", level=1)

    steps = [
        ("1", "Device 선택 (↺ 새로고침으로 단말 목록 갱신)"),
        ("2", "저장 폴더, 사이클 수 등 파라미터 설정"),
        ("3", "테스트 단계 확인 / 편집"),
        ("4", "정상 조건 확인 / 편집"),
        ("5", "[▶ START] 클릭"),
        ("6", "[강제 Crash 옵션 ON인 경우]\n"
              "  6-1. AT!ERROROPT? 확인\n"
              "  6-2. USB Dump 미포함 → AT!ERROROPT=0 + AT+CFUN=1,1 재부팅\n"
              "  6-3. adb devices 로 단말 복귀 확인 (최대 120초)\n"
              "  6-4. 단말 미복귀 시 테스트 중단"),
        ("7", "QXDM이 자동으로 실행됨"),
        ("8", "QXDM에서 Phone 연결 (Ctrl+O)"),
        ("9", "[✔ Phone 연결 완료] 클릭 → 테스트 자동 진행"),
        ("10", "비정상 감지 시 ISF 자동 저장 후 QXDM 종료"),
    ]
    _add_table(doc,
        headers=["단계", "내용"],
        rows=steps,
        col_widths=[1.5, 13.5],
    )
    _add_note(doc, "[■ STOP] 버튼으로 테스트 중간에 수동 중단 가능", bg="DEEBF7")

    doc.add_heading("전체 실행 흐름도", level=2)
    _add_flow_image(doc, _flow_image_test(), "테스트 실행 흐름", width=Inches(3.2))

    doc.add_heading("강제 Crash 옵션 흐름도", level=2)
    _add_flow_image(doc, _flow_image_crash(), "강제 Crash 옵션 동작", width=Inches(4.2))


def _section_logfiles(doc):
    doc.add_heading("8. 로그 파일", level=1)
    p = doc.add_paragraph("테스트 시작 시 저장 폴더 내에 자동 생성됩니다.")
    for run in p.runs:
        run.font.name = "맑은 고딕"; run.font.size = Pt(10)
    _add_table(doc,
        headers=["파일명 패턴", "내용"],
        rows=[
            ["test_YYYYMMDD_HHMMSS.log",            "전체 실행 텍스트 로그"],
            ["test_..._FINAL_....isf",               "정상 종료 시 최종 ISF"],
            ["test_..._TRIGGER_cycle###_....isf",    "비정상 감지 시 ISF"],
            ["test_..._part###_....isf",             "롤링 저장 중간 ISF"],
        ],
        col_widths=[7, 8],
    )


def _section_rolling(doc):
    doc.add_heading("9. 롤링 저장", level=1)
    p = doc.add_paragraph("장시간 테스트 시 QXDM 메모리 누적을 방지하기 위한 기능입니다.")
    for run in p.runs:
        run.font.name = "맑은 고딕"; run.font.size = Pt(10)
    items = [
        "롤링 저장 주기 = N : N 사이클마다 ISF 저장 후 QXDM 로그 클리어",
        "롤링 보관 개수 = M : 최근 M개 파일만 유지, 초과분 자동 삭제",
        "최종 비정상 감지 시에는 별도로 TRIGGER ISF가 추가 저장됨",
    ]
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "맑은 고딕"; run.font.size = Pt(10)
    doc.add_paragraph()


def _section_caution(doc):
    doc.add_heading("10. 주의 사항", level=1)
    items = [
        "QXDM이 이미 실행 중이면 COM 충돌이 발생할 수 있습니다. 테스트 전 QXDM을 종료해 주세요.",
        "ADB가 PATH에 없으면 모든 명령이 실패합니다. adb devices 명령으로 사전 확인하세요.",
        "강제 Crash 옵션은 단말을 강제로 재부팅시킵니다. 필요한 경우에만 사용하세요.",
        "Windows 종료 옵션 사용 시 종료 대기 시간 내에 shutdown /a 명령으로 취소 가능합니다.",
        "FUNCTION 조건에서 adb_shell 호출 시 오류 메시지가 출력에 포함될 수 있으므로 2>/dev/null을 사용하세요.",
    ]
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "맑은 고딕"; run.font.size = Pt(10)
    doc.add_paragraph()


def _section_version(doc):
    doc.add_heading("11. 버전 관리", level=1)
    p = doc.add_paragraph(
        "버전 정보는 clat_test_gui.py 상단 두 줄로 관리합니다. "
        "이 값을 변경하면 타이틀바와 EXE 파일명이 자동으로 반영됩니다."
    )
    for run in p.runs:
        run.font.name = "맑은 고딕"; run.font.size = Pt(10)
    _add_code(doc,
        '__version__ = "1.0.0"   # ← 버전 수정 시 이 값만 변경\n'
        'APP_NAME    = "QXDM AUTO LOG"'
    )
    _add_table(doc,
        headers=["항목", "형식"],
        rows=[
            ["타이틀바",    "QXDM AUTO LOG  v1.0.0"],
            ["EXE 파일명",  "qxdm_auto_log_v1.0.0.exe"],
        ],
        col_widths=[4, 11],
    )


def _section_build(doc):
    doc.add_heading("12. 재빌드 방법", level=1)
    p = doc.add_paragraph(
        "clat_test_gui.py 수정 후 EXE 재생성이 필요할 때 build_exe.bat 을 실행합니다. "
        "__version__ 값을 자동으로 읽어 EXE 이름에 반영합니다."
    )
    for run in p.runs:
        run.font.name = "맑은 고딕"; run.font.size = Pt(10)
    _add_code(doc, "build_exe.bat  더블클릭")
    p2 = doc.add_paragraph("결과물: ")
    r = p2.add_run("dist\\qxdm_auto_log_v{버전}.exe")
    r.font.name = "Consolas"; r.font.size = Pt(10)
    for run in p2.runs:
        run.font.name = "맑은 고딕"; run.font.size = Pt(10)
    p2.runs[0].font.name = "맑은 고딕"
    doc.add_paragraph()


def _section_changelog(doc):
    doc.add_heading("변경 이력", level=1)
    _add_table(doc,
        headers=["버전", "변경 내용"],
        rows=[
            ["v1.0.0",
             "최초 릴리즈\n"
             "· 버전 정보 표시 (타이틀바, EXE 파일명)\n"
             "· 강제 Crash 옵션 동작 변경: 테스트 시작 전 AT!ERROROPT 확인 및 재부팅\n"
             "· 조건 편집 (SHELL / AT / FUNCTION) 지원\n"
             "· 테스트 단계 편집 지원\n"
             "· 롤링 ISF 저장 지원\n"
             "· cmd 창 팝업 억제"],
        ],
        col_widths=[2.5, 12.5],
    )


# ─────────────────────────────────────────────
# 메인
# ─────────────────────────────────────────────
def main():
    print(f"[*] 매뉴얼 생성 중: {OUT_FILE}")

    doc = Document()
    _apply_default_styles(doc)

    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

    _title_page(doc)
    _section_overview(doc)
    _section_requirements(doc)
    _section_screen(doc)
    _section_params(doc)
    _section_steps(doc)
    _section_conditions(doc)
    _section_procedure(doc)
    _section_logfiles(doc)
    _section_rolling(doc)
    _section_caution(doc)
    _section_version(doc)
    _section_build(doc)
    _section_changelog(doc)

    doc.save(OUT_FILE)
    print(f"[완료] {OUT_FILE}")


if __name__ == "__main__":
    main()
